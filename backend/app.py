import os
import json
import threading
import re
import openai
import PyPDF2
from fpdf import FPDF
from datetime import datetime
import copy
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import glob
import time

# --- GESTION DES DÉPENDANCES OPTIONNELLES ---
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("AVERTISSEMENT: python-docx non installé. L'export Word ne sera pas disponible.")

# --- CONFIGURATION DE L'API OPENAI ---
# IMPORTANT: Définissez cette variable d'environnement sur votre serveur !
# Exemple: export OPENAI_API_KEY='sk-...'
openai.api_key = os.getenv("OPENAI_API_KEY")

# --- INITIALISATION DE FLASK ---
app = Flask(__name__)
# CORS permet à votre frontend (ex: http://localhost:5173) de communiquer avec ce backend (http://127.0.0.1:5000)
CORS(app) 

# --- CONSTANTES DE L'APPLICATION (reprises de votre script) ---
PLACEHOLDER_PATTERN = r'(\[(?:À PRÉCISER|À VALIDER|COMPLÉTER|RÉFÉRENCE À INDIQUER).*?\])'
REFERENCE_PATTERN = r'(\(voir exemple CCTP .*?(?:\s*->\s*.*?)?\))'
CROSS_REF_PATTERN = r'(\{\{REF:(.*?)\|(.*?)\}\})'
MAX_EXAMPLES_IN_PROMPT = 3
MAX_TOTAL_EXAMPLE_CHARS = 4000

# Variable globale pour le statut de l'analyse
analysis_status = {"running": False, "progress": 0, "max_files": 0, "current_file": "", "error": None}

# --- CHEMINS VERS LES DOSSIERS DE DONNÉES ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODELES_DIR = os.path.join(BASE_DIR, "modeles_cctp")
PREVIEWS_DIR = os.path.join(BASE_DIR, "previews_cctp")
KNOWLEDGE_BASE_DIR = os.path.join(BASE_DIR, "knowledge_base")
KNOWLEDGE_BASE_PATH = os.path.join(KNOWLEDGE_BASE_DIR, "knowledge_base.txt")
BIBLIOTHEQUE_SECTIONS_PATH = os.path.join(BASE_DIR, "bibliotheque_sections.json")
SYSTEM_PROMPT_PATH = os.path.join(BASE_DIR, "system_prompt.txt")

# S'assurer que les dossiers de données existent au démarrage
os.makedirs(MODELES_DIR, exist_ok=True)
os.makedirs(PREVIEWS_DIR, exist_ok=True)
os.makedirs(KNOWLEDGE_BASE_DIR, exist_ok=True)

# --- PROMPT PAR DÉFAUT ---
DEFAULT_SYSTEM_PROMPT = (
    "Tu es un expert rédacteur de CCTP pour des projets de construction (phase PRO). Ta mission est de rédiger une description technique précise pour une section spécifique à partir des notes (contexte) qui te sont données.\n\n"
    "INSTRUCTIONS IMPÉRATIVES :\n"
    "1.  **Niveau de détail 'PRO'** : La précision est essentielle. Évite les généralités. Sois technique et factuel. MAIS SURTOUT, adapte le niveau de détail à la complexité de la section (ex : 'Localisation' = bref, 'Performances' = très détaillé). C'EST TRÈS IMPORTANT D'ÊTRE PRÉCIS MAIS SANS TROP DONNER D'ÉLÉMENTS, PRÉCIS MAIS COURT !\n"
    "2.  **Placeholders** : S'il manque des informations essentielles pour un niveau PRO (valeurs, références, choix techniques, etc.), insère un placeholder clair. Ex : `[À PRÉCISER : type de vitrage]`, `[À VALIDER : résistance au feu EI60]`. Il vaut MIEUX signaler ce qui manque que de le deviner. Il faut que ce soit précis, clair, équilibré et pertinent.\n"
    "3.  **Utilisation des exemples fournis (CONSEILLÉE)** : Tu disposes d'extraits de CCTP fournis en exemple. **Essaie d'en réutiliser un ou plusieurs quand c'est pertinent**, pour enrichir ou formuler plus clairement certains éléments techniques. **Ce n'est pas obligatoire**, mais c'est **fortement recommandé** quand un exemple est utile, crédible ou mieux formulé que ce que tu peux déduire uniquement des notes. **Ne dépasse pas 1 à 2 utilisations par section.** Si tu reprends un extrait ou t'en inspires, mentionne la source avec le format suivant : `(voir exemple CCTP nom_du_fichier.pdf -> nom_section_originale)` où nom_section_originale est le nom exact de la section dans le document source.\n"
    "4.  **Formatage** : Réponds DIRECTEMENT avec le texte de la section. N'inclus PAS de titre, de numérotation ni le nom de la typologie dans ta réponse."
)

def load_system_prompt():
    """Charge le prompt système personnalisé, ou retourne le prompt par défaut si absent."""
    if os.path.exists(SYSTEM_PROMPT_PATH):
        try:
            with open(SYSTEM_PROMPT_PATH, "r", encoding="utf-8") as f:
                prompt = f.read().strip()
                if prompt:
                    return prompt
        except Exception:
            pass
    return DEFAULT_SYSTEM_PROMPT

def save_system_prompt(new_prompt):
    """Sauvegarde le prompt système personnalisé."""
    with open(SYSTEM_PROMPT_PATH, "w", encoding="utf-8") as f:
        f.write(new_prompt.strip())

@app.route('/api/system-prompt', methods=['GET'])
def get_system_prompt():
    """Renvoie le prompt système courant (personnalisé ou défaut)."""
    return jsonify({"prompt": load_system_prompt()})

@app.route('/api/system-prompt', methods=['POST'])
def set_system_prompt():
    """Sauvegarde un nouveau prompt système personnalisé."""
    data = request.json
    prompt = data.get("prompt", "").strip()
    if not prompt:
        return jsonify({"error": "Le prompt ne peut pas être vide."}), 400
    save_system_prompt(prompt)
    return jsonify({"message": "Prompt système sauvegardé."})

# --- CLASSE PDF (identique à votre script original) ---
class PDF(FPDF):
    def __init__(self, *args, **kwargs):
        self.chapter_map = kwargs.pop('chapter_map', [])
        super().__init__(*args, **kwargs)
        try:
            # IMPORTANT: Assurez-vous que le dossier `fonts` se trouve dans `backend/`
            self.add_font('DejaVu', '', 'fonts/DejaVuSans.ttf', uni=True)
            self.add_font('DejaVu', 'B', 'fonts/DejaVuSans-Bold.ttf', uni=True)
        except RuntimeError as e:
            print(f"ERREUR FATALE: Impossible de trouver les polices pour FPDF. {e}")
            print("Assurez-vous que les fichiers 'DejaVuSans.ttf' et 'DejaVuSans-Bold.ttf' sont dans un dossier nommé 'fonts' à la racine du dossier 'backend'.")
            raise SystemExit("Polices FPDF non trouvées.")

        self.normal_color = (105, 108, 111)
        self.placeholder_color = (220, 53, 69)
        self.reference_color = (40, 167, 69)
        self.cross_ref_color = (0, 123, 255)
        self.set_text_color(*self.normal_color)

    def header(self):
        self.set_font('DejaVu', 'B', 9)
        self.set_text_color(180, 180, 180)
        self.cell(0, 8, 'Description générale des ouvrages', border=0, align='L', ln=1)
        self.ln(3)
        self.set_text_color(*self.normal_color)

    def footer(self):
        self.set_y(-15)
        self.set_font('DejaVu', '', 8)
        self.set_text_color(*self.normal_color)
        self.cell(0, 8, f'Page {self.page_no()}/{{nb}}', align='C')

    def add_typology_title(self, typo_index, title):
        self.set_font('DejaVu', 'B', 12)
        self.set_text_color(0, 0, 0)
        self.cell(0, 10, f"{typo_index + 1}. {title}", border=0, ln=1, align='L')
        self.ln(2)
        self.set_font('DejaVu', '', 9)
        self.set_text_color(*self.normal_color)

    def add_section_title(self, number, title):
        self.set_font('DejaVu', 'B', 9)
        self.set_text_color(*self.normal_color)
        self.set_x(15)
        self.cell(0, 8, f'{number} {title}', border='B', ln=1, align='L')
        self.ln(1)

    def add_body_text(self, text):
        current_x = self.get_x()
        self.set_x(20)
        # Nettoyer le texte des marqueurs de formatage
        text_cleaned = re.sub(r'^###.*$', '', text, flags=re.MULTILINE)
        text_cleaned = text_cleaned.replace('•', '-').strip()
        # Écrire tout en couleur normale, sans distinction
        self.set_font('DejaVu', '', 9)
        self.set_text_color(*self.normal_color)
        self.multi_cell(0, 5, text_cleaned)
        self.ln(5)
        self.set_x(current_x)

    def _render_cross_references(self, text):
        """Cette méthode n'est plus utilisée car les transformations sont faites en amont"""
        return text

def extract_parts_for_export(text, chapter_map, include_ai_notes=True):
    """Sépare texte principal, placeholders, exemples, crossrefs pour export."""
    if not text:
        return "", [], [], []
    
    # Nettoyer le texte d'abord
    cleaned_text = _clean_text_for_export(text)
    
    placeholders = []
    exemples = []
    crossrefs = []
    main_text = cleaned_text
    
    # Extraire les placeholders [À ...] si demandé
    if include_ai_notes:
        placeholder_matches = re.findall(PLACEHOLDER_PATTERN, main_text)
        for match in placeholder_matches:
            if match not in placeholders:  # Éviter les doublons
                placeholders.append(match)
    main_text = re.sub(PLACEHOLDER_PATTERN, '', main_text)
    
    # Extraire les exemples (voir exemple CCTP ...) si demandé
    if include_ai_notes:
        example_matches = re.findall(REFERENCE_PATTERN, main_text)
        for match in example_matches:
            if match not in exemples:  # Éviter les doublons
                exemples.append(match)
    main_text = re.sub(REFERENCE_PATTERN, '', main_text)
    
    # Traiter les cross-references {{REF:...|...}}
    def crossref_repl(match):
        typo = match.group(1)
        section = match.group(2)
        ref = next((c for c in chapter_map if c['nom_typo'] == typo and c['titre_section'] == section), None)
        if ref:
            ref_text = f"(voir section {ref['number']} {section})"
            if ref_text not in crossrefs:  # Éviter les doublons
                crossrefs.append(ref_text)
        else:
            error_ref = f"(référence introuvable: {typo}|{section})"
            if error_ref not in crossrefs:  # Éviter les doublons
                crossrefs.append(error_ref)
        return ''
    
    main_text = re.sub(CROSS_REF_PATTERN, crossref_repl, main_text)
    
    # Nettoyer le texte principal des espaces multiples
    main_text = re.sub(r'\s+', ' ', main_text).strip()
    
    return main_text, placeholders, exemples, crossrefs

# --- FONCTIONS LOGIQUES (adaptées de votre script) ---

def openai_chat_completion(model, messages, temperature=0.3):
    """Wrapper compatible avec openai>=1.0."""
    return openai.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
    )

def _get_base_prompt_generation():
    """Fonction qui contient le prompt de base de VF_LOW_TOKEN.py"""
    return (
        "Tu es un expert rédacteur de CCTP pour des projets de construction (phase PRO). Ta mission est de rédiger une description technique précise pour une section spécifique à partir des notes (contexte) qui te sont données.\n\n"
        "INSTRUCTIONS IMPÉRATIVES :\n"
        "1.  **Niveau de détail 'PRO'** : La précision est essentielle. Évite les généralités. Sois technique et factuel. MAIS SURTOUT, adapte le niveau de détail à la complexité de la section (ex : 'Localisation' = bref, 'Performances' = très détaillé). C'EST TRÈS IMPORTANT D'ÊTRE PRÉCIS MAIS SANS TROP DONNER D'ÉLÉMENTS, PRÉCIS MAIS COURT !\n"
        "2.  **Placeholders** : S'il manque des informations essentielles pour un niveau PRO (valeurs, références, choix techniques, etc.), insère un placeholder clair. Ex : `[À PRÉCISER : type de vitrage]`, `[À VALIDER : résistance au feu EI60]`. Il vaut MIEUX signaler ce qui manque que de le deviner. Il faut que ce soit précis, clair, équilibré et pertinent.\n"
        "3.  **Utilisation des exemples fournis (CONSEILLÉE)** : Tu disposes d'extraits de CCTP fournis en exemple. **Essaie d'en réutiliser un ou plusieurs quand c'est pertinent**, pour enrichir ou formuler plus clairement certains éléments techniques. **Ce n'est pas obligatoire**, mais c'est **fortement recommandé** quand un exemple est utile, crédible ou mieux formulé que ce que tu peux déduire uniquement des notes. **Ne dépasse pas 1 à 2 utilisations par section.** Si tu reprends un extrait ou t'en inspires, mentionne la source avec le format suivant : `(voir exemple CCTP nom_du_fichier.pdf -> nom_section_originale)` où nom_section_originale est le nom exact de la section dans le document source.\n"
        "4.  **Formatage** : Réponds DIRECTEMENT avec le texte de la section. N'inclus PAS de titre, de numérotation ni le nom de la typologie dans ta réponse."
    )

def _apply_prompt_modifications(base_prompt, modification_text):
    """Ajoute les modifications comme instructions supplémentaires au prompt de base"""
    if not modification_text or not modification_text.strip():
        return base_prompt
    
    try:
        # Demander à GPT de transformer la modification en instruction claire à ajouter au prompt
        messages = [
            {"role": "system", "content": "Tu es un expert en rédaction de prompts pour IA. Tu dois transformer une demande utilisateur en instruction claire à ajouter à un prompt existant."},
            {"role": "user", "content": f"""Voici le prompt de base pour générer des textes de CCTP :
---
{base_prompt}
---

L'utilisateur souhaite ajouter cette modification/contrainte :
"{modification_text}"

Ta tâche : Transforme cette demande en instruction claire et précise à ajouter à la fin du prompt, sous forme d'une nouvelle instruction numérotée (5.) qui s'intègre harmonieusement avec les 4 instructions existantes.

L'instruction doit :
- Être claire et actionnable
- Être formulée comme une consigne pour l'IA
- S'intégrer naturellement avec le style du prompt existant
- Être numérotée "5." pour suivre la logique

Réponds uniquement avec le prompt complet incluant la nouvelle instruction."""}
        ]
        
        response = openai_chat_completion(
            model="gpt-4-1106-preview",
            messages=messages,
            temperature=0.3,
        )
        
        modified_prompt = response.choices[0].message.content.strip()
        return modified_prompt
        
    except Exception as e:
        print(f"Erreur lors de l'ajout des instructions au prompt: {e}")
        # En cas d'erreur, ajouter simplement l'instruction à la fin
        return base_prompt + f"\n\n5.  **Instruction supplémentaire** : {modification_text}"

def _construire_prompt_generation(nom_typo, titre_section_propre, notes_utilisateur, contexte_precedent, exemples_pertinents, custom_instruction):
    # Obtenir le prompt de base de VF_LOW_TOKEN
    base_instruction = _get_base_prompt_generation()
    
    # Obtenir le prompt système personnalisé s'il existe
    system_prompt = load_system_prompt()
    
    # Si le prompt système n'est pas le prompt par défaut, cela signifie qu'il a été modifié
    if system_prompt != DEFAULT_SYSTEM_PROMPT:
        # Utiliser le prompt système personnalisé comme instruction de base
        instruction_ia = system_prompt
    else:
        # Utiliser le prompt de base original
        instruction_ia = base_instruction
    
    # Ajouter la consigne spécifique pour cette section si elle existe
    if custom_instruction:
        instruction_ia += f"\n\nCONSIGNE SPÉCIFIQUE POUR CETTE SECTION: {custom_instruction}"

    # Construire le prompt final avec les variables
    base_prompt = (
        f"{instruction_ia}\n\n"
        f"---\n\n"
        f"**Typologie à décrire**: \"{nom_typo}\"\n"
        f"**Section à rédiger**: \"{titre_section_propre}\"\n\n"
        f"**Notes de l'utilisateur pour cette section**:\n```\n{notes_utilisateur or 'Aucune note fournie.'}\n```\n\n"
        f"**Contexte (structure des autres sections de la typologie)**:\n{contexte_precedent or 'Aucun.'}\n\n"
        "**Exemples pertinents (extraits de CCTP existants)**:\n"
        f"{_format_exemples_pour_prompt(exemples_pertinents) if exemples_pertinents else 'Aucun.'}\n"
    )
    
    return base_prompt

def _construire_prompt_modification(action, nom_typo, titre_section_propre, texte_actuel_ia):
    action_map = {
        "correct": "Corrige le texte suivant pour la grammaire, l'orthographe, et la syntaxe. Ne modifie PAS le contenu technique, les valeurs ou les placeholders.",
        "lengthen": "Développe le texte suivant en ajoutant des détails techniques plausibles de niveau PRO.",
        "shorten": "Résume le texte suivant pour qu'il soit plus concis."
    }
    return (
        "MISSION: Tu es un assistant expert en rédaction de CCTP. Modifie un texte existant selon l'instruction.\n\n"
        f"INSTRUCTION: {action_map.get(action, 'Modifie le texte')}\n\n"
        f"TYPOLOGIE: \"{nom_typo}\"\n"
        f"SECTION: \"{titre_section_propre}\"\n\n"
        "TEXTE ORIGINAL:\n---\n"
        f"{texte_actuel_ia}\n---\n\n"
        "Ta réponse doit contenir UNIQUEMENT le texte modifié, sans titre ni introduction."
    )

def _retrouver_exemples_pertinents(titre_section_cible, contenu_kb_json):
    """Recherche des exemples pertinents dans la knowledge base JSON avec le nom de section d'origine"""
    if not contenu_kb_json:
        return []
    
    exemples_trouves = []
    try:
        # Si c'est un fichier JSON, on le parse
        if isinstance(contenu_kb_json, str):
            try:
                kb_data = json.loads(contenu_kb_json)
            except json.JSONDecodeError:
                # Si ce n'est pas du JSON, on utilise l'ancienne méthode
                return _retrouver_exemples_pertinents_legacy(titre_section_cible, contenu_kb_json)
        else:
            kb_data = contenu_kb_json
        
        # Parcourir tous les fichiers dans la knowledge base
        for filename, typologies in kb_data.items():
            if not isinstance(typologies, list):
                continue
                
            for typologie in typologies:
                if not isinstance(typologie, dict) or 'sections' not in typologie:
                    continue
                    
                for section in typologie['sections']:
                    if not isinstance(section, dict):
                        continue
                        
                    titre_section_original = section.get('titre', '').strip()
                    contenu_section = section.get('contenu', '').strip()
                    
                    # Recherche par correspondance exacte ou partielle
                    if (titre_section_original.lower() == titre_section_cible.lower() or 
                        titre_section_cible.lower() in titre_section_original.lower() or
                        titre_section_original.lower() in titre_section_cible.lower()):
                        
                        if contenu_section:
                            exemples_trouves.append({
                                "section": titre_section_cible,
                                "section_originale": titre_section_original,  # Nouveau champ
                                "source": filename,
                                "texte": contenu_section
                            })
    except Exception as e:
        print(f"Erreur lors de la récupération des exemples: {e}")
        return []
    
    # Limiter les exemples
    limited_examples = []
    total_chars = 0
    for ex in exemples_trouves:
        if len(limited_examples) >= MAX_EXAMPLES_IN_PROMPT:
            break
        if total_chars + len(ex["texte"]) > MAX_TOTAL_EXAMPLE_CHARS:
            chars_to_add = MAX_TOTAL_EXAMPLE_CHARS - total_chars
            if chars_to_add > 100:
                ex["texte"] = ex["texte"][:chars_to_add] + "..."
                limited_examples.append(ex)
            break
        limited_examples.append(ex)
        total_chars += len(ex["texte"])
    
    return limited_examples

def _retrouver_exemples_pertinents_legacy(titre_section_cible, contenu_total_exemples):
    """Méthode de fallback pour l'ancien format de knowledge base"""
    if not contenu_total_exemples: 
        return []
    
    exemples_trouves = []
    try:
        document_chunks = contenu_total_exemples.split('--- EXTRAIT DU DOCUMENT : ')[1:]
        for chunk in document_chunks:
            try:
                lines = chunk.split('\n', 1)
                if len(lines) < 2:
                    continue
                filename = lines[0].replace('---', '').strip()
                document_content = lines[1] if len(lines) > 1 else ""
                
                # Rechercher la section dans le contenu
                pattern = re.compile(r"^\s*(?:\d+\.\d+\s+|###\s+)? *" + re.escape(titre_section_cible) + r"\s*$(.*?)(?=^\s*(?:\d+\.\d+\s+|###\s+)?\w|\Z)", re.IGNORECASE | re.MULTILINE | re.DOTALL)
                for match in pattern.finditer(document_content):
                    contenu_section = match.group(1).strip()
                    if contenu_section: 
                        exemples_trouves.append({
                            "section": titre_section_cible,
                            "section_originale": titre_section_cible,  # Même nom pour legacy
                            "source": filename, 
                            "texte": contenu_section
                        })
            except Exception as e:
                print(f"Avertissement: Erreur de parsing d'un chunk de la KB: {e}")
                continue
    except Exception as e:
        print(f"Erreur lors de la récupération des exemples: {e}")
        return []
    
    # Limiter les exemples
    limited_examples = []
    total_chars = 0
    for ex in exemples_trouves:
        if len(limited_examples) >= MAX_EXAMPLES_IN_PROMPT:
            break
        if total_chars + len(ex["texte"]) > MAX_TOTAL_EXAMPLE_CHARS:
            chars_to_add = MAX_TOTAL_EXAMPLE_CHARS - total_chars
            if chars_to_add > 100:
                ex["texte"] = ex["texte"][:chars_to_add] + "..."
                limited_examples.append(ex)
            break
        limited_examples.append(ex)
        total_chars += len(ex["texte"])
    
    return limited_examples

def _format_exemples_pour_prompt(exemples):
    """Formate les exemples pour le prompt."""
    if not exemples:
        return "Aucun exemple disponible."
    
    formatted = ""
    for i, exemple in enumerate(exemples, 1):
        section_originale = exemple.get('section_originale', exemple.get('section', ''))
        source = exemple.get('source', 'source inconnue')
        texte = exemple.get('texte', '')
        
        formatted += f"**Exemple {i}** (source: {source} -> {section_originale}):\n{texte}\n\n"
    
    return formatted.strip()

def enrichir_bibliotheque_sections():
    """Met à jour la bibliothèque de sections à partir des modèles existants."""
    sections = set()
    
    try:
        # Parcourir tous les modèles
        for filename in os.listdir(MODELES_DIR):
            if filename.endswith('.json'):
                file_path = os.path.join(MODELES_DIR, filename)
                with open(file_path, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                
                # Extraire les titres de sections
                for typologie in project_data:
                    for section in typologie.get('sections', []):
                        titre = section.get('titre', '').strip()
                        if titre:
                            sections.add(titre)
        
        # Sauvegarder la bibliothèque mise à jour
        sections_list = sorted(list(sections))
        with open(BIBLIOTHEQUE_SECTIONS_PATH, 'w', encoding='utf-8') as f:
            json.dump(sections_list, f, indent=2, ensure_ascii=False)
            
    except Exception as e:
        print(f"Erreur lors de l'enrichissement de la bibliothèque: {e}")
        # Créer une bibliothèque vide si erreur
        with open(BIBLIOTHEQUE_SECTIONS_PATH, 'w', encoding='utf-8') as f:
            json.dump([], f)

def _render_cross_references_for_export(text, chapter_map):
    """Rend les références croisées pour l'export."""
    if not text:
        return ""
    
    def replace_func(match):
        nom_typo = match.group(1)
        nom_section = match.group(2)
        
        for item in chapter_map:
            if item['nom_typo'] == nom_typo and item['titre_section'] == nom_section:
                return f"(voir section {item['number']} {nom_section})"
        return f"(référence introuvable: {nom_typo}|{nom_section})"
    
    # Transformer uniquement les références non encore transformées
    return re.sub(CROSS_REF_PATTERN, replace_func, text)

def _clean_text_for_export(text):
    """Nettoie le texte pour l'export en évitant les doublons."""
    if not text:
        return ""
    
    # Supprimer toutes les balises HTML
    text = re.sub(r'<[^>]+>', '', text)
    
    # Nettoyer les entités HTML
    text = text.replace('&lt;', '<').replace('&gt;', '>')
    
    # Nettoyer les espaces multiples et les retours à la ligne multiples
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

# --- ROUTES DE L'API (les points d'accès pour le frontend) ---
def get_kb_status():
    """Renvoie le statut de la base de connaissance."""
    if os.path.exists(KNOWLEDGE_BASE_PATH):
        mod_time = os.path.getmtime(KNOWLEDGE_BASE_PATH)
        return f"Base analysée le {datetime.fromtimestamp(mod_time).strftime('%d/%m/%Y %H:%M')}"
    else:
        return "Base de connaissance non trouvée."
    
@app.route('/api/status', methods=['GET'])
def get_status():
    """Route simple pour vérifier que le serveur est en ligne et si l'API key est configurée."""
    return jsonify({
        "status": "online",
        "openai_configured": bool(openai.api_key),
        "docx_available": DOCX_AVAILABLE,
        "knowledge_base_status": get_kb_status()
    })

@app.route('/api/models', methods=['GET'])
def get_models():
    """Liste les modèles de projet (.json) disponibles dans le dossier modeles_cctp."""
    try:
        models = sorted([f[:-5] for f in os.listdir(MODELES_DIR) if f.endswith(".json")])
        return jsonify(models)
    except Exception as e:
        return jsonify({"error": f"Impossible de lister les modèles: {str(e)}"}), 500

@app.route('/api/models', methods=['POST'])
def create_model():
    """Crée un nouveau fichier de modèle vide."""
    data = request.json
    model_name = data.get('name')
    if not model_name or not model_name.strip():
        return jsonify({"error": "Le nom du modèle est requis."}), 400
    
    file_path = os.path.join(MODELES_DIR, f"{model_name.strip()}.json")
    if os.path.exists(file_path):
        return jsonify({"error": "Ce nom de modèle existe déjà."}), 409

    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump([], f) # Crée un projet avec une liste de typologies vide
        return jsonify({"message": f"Modèle '{model_name}' créé avec succès."}), 201
    except Exception as e:
        return jsonify({"error": f"Impossible de créer le fichier modèle: {str(e)}"}), 500

@app.route('/api/data/<model_name>', methods=['GET'])
def get_model_data(model_name):
    """Charge les données d'un modèle (projet + prévisualisations)."""
    project_path = os.path.join(MODELES_DIR, f"{model_name}.json")
    previews_path = os.path.join(PREVIEWS_DIR, f"{model_name}_previews.json")

    if not os.path.exists(project_path):
        return jsonify({"error": "Modèle non trouvé."}), 404

    try:
        with open(project_path, 'r', encoding='utf-8') as f:
            project_data = json.load(f)
        
        previews_data = {}
        if os.path.exists(previews_path):
            with open(previews_path, 'r', encoding='utf-8') as f:
                previews_data = json.load(f)
        
        with open(BIBLIOTHEQUE_SECTIONS_PATH, 'r', encoding='utf-8') as f:
            sections_library = json.load(f)

        return jsonify({"project": project_data, "previews": previews_data, "sectionsLibrary": sections_library})
    except Exception as e:
        return jsonify({"error": f"Erreur de lecture des fichiers du modèle: {str(e)}"}), 500

@app.route('/api/data/<model_name>', methods=['POST'])
def save_model_data(model_name):
    """Sauvegarde les données (projet + prévisualisations) pour un modèle donné."""
    data = request.json
    project_data = data.get('project')
    previews_data = data.get('previews')

    project_path = os.path.join(MODELES_DIR, f"{model_name}.json")
    previews_path = os.path.join(PREVIEWS_DIR, f"{model_name}_previews.json")

    try:
        with open(project_path, 'w', encoding='utf-8') as f:
            json.dump(project_data, f, indent=2, ensure_ascii=False)
        with open(previews_path, 'w', encoding='utf-8') as f:
            json.dump(previews_data, f, indent=2, ensure_ascii=False)
        
        # Mettre à jour la bibliothèque avec les nouvelles sections créées
        enrichir_bibliotheque_sections()
        
        return jsonify({"message": "Projet sauvegardé avec succès."})
    except Exception as e:
        return jsonify({"error": f"Erreur lors de la sauvegarde du projet: {str(e)}"}), 500

@app.route('/api/generate', methods=['POST'])
def handle_generation():
    """Point d'accès principal pour la génération de texte par l'IA."""
    if not openai.api_key:
        return jsonify({"error": "La clé API OpenAI n'est pas configurée sur le serveur."}), 503
    
    data = request.json
    try:
        # Extraire les données nécessaires de la requête
        nom_typo = data['nomTypo']
        titre_section = data['titreSection']
        notes_utilisateur = data.get('notes', '')
        texte_actuel_ia = data.get('texteActuel', '')
        action = data.get('action', 'génération du contenu...')
        contexte_summarized = data.get('contexteSummarized', '')
        custom_instruction = data.get('customInstruction', '')

        contenu_kb = {}
        # Charger la knowledge base JSON
        kb_json_path = os.path.join(KNOWLEDGE_BASE_DIR, "knowledge_base.json")
        if os.path.exists(kb_json_path):
            with open(kb_json_path, "r", encoding="utf-8") as f:
                contenu_kb = json.load(f)
        # Fallback vers le format texte si JSON n'existe pas
        elif os.path.exists(KNOWLEDGE_BASE_PATH):
            with open(KNOWLEDGE_BASE_PATH, "r", encoding="utf-8") as f:
                contenu_kb = f.read()
        
        exemples = _retrouver_exemples_pertinents(titre_section, contenu_kb)

        if action == "génération du contenu...":
            prompt = _construire_prompt_generation(nom_typo, titre_section, notes_utilisateur, contexte_summarized, exemples, custom_instruction)
            model_to_use = "gpt-4-1106-preview"
        else:
            prompt = _construire_prompt_modification(action, nom_typo, titre_section, texte_actuel_ia)
            model_to_use = "gpt-3.5-turbo-0125"

        response = openai_chat_completion(
            model=model_to_use,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        texte_genere = response.choices[0].message.content.strip()

        # --- Nettoyage du texte généré (robuste) ---
        def clean_ai_output(text):
            try:
                lines = text.splitlines()
                cleaned = []
                for line in lines:
                    l = line.strip()
                    # Supprimer les séparateurs "---"
                    if l == "---":
                        continue
                    # Supprimer les titres en gras markdown (**TITRE**)
                    if re.match(r"^\*{2}.+\*{2}$", l):
                        continue
                    # Supprimer les titres tout en majuscules (hors phrases normales)
                    if l.isupper() and len(l) < 80:
                        continue
                    # Supprimer les titres type markdown (# ou ## ...)
                    if re.match(r"^#+\s", l):
                        continue
                    # Supprimer les lignes vides
                    if not l:
                        continue
                    cleaned.append(line)
                return "\n".join(cleaned).strip()
            except Exception as e:
                print(f"[CCTP] Erreur nettoyage texte IA: {e}")
                return text.strip() if text else ""

        texte_genere = clean_ai_output(texte_genere)
        # Si le texte est vide, retourne une chaîne vide proprement
        if texte_genere is None:
            texte_genere = ""
        return jsonify({"text": texte_genere})

    except KeyError as e:
        print(f"[CCTP] Erreur KeyError: {e}")
        return jsonify({"error": f"Donnée manquante dans la requête: {e}"}), 400
    except Exception as e:
        print(f"[CCTP] Erreur lors de la génération OpenAI: {e}")
        return jsonify({"error": f"Erreur lors de la génération OpenAI: {str(e)}"}), 500

@app.route('/api/export/pdf', methods=['POST'])
def export_pdf():
    """Génère un fichier PDF à partir des données du projet et le renvoie pour téléchargement."""
    data = request.json
    project_data = data.get('project', [])
    previews_data = data.get('previews', {})
    chapter_map = data.get('chapterMap', [])
    include_ai_notes = data.get('includeAiNotes', True)  # Par défaut, inclure les notes IA

    try:
        pdf_path = os.path.join(BASE_DIR, 'cctp_export.pdf')
        pdf = PDF('P', 'mm', 'A4', chapter_map=chapter_map)
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        for typo_idx, typo_data in enumerate(project_data):
            nom_typo = typo_data.get("nomTypologie")
            typo_previews = previews_data.get(nom_typo, {})
            if not any(typo_previews.values()): 
                continue
            if pdf.get_y() > 230 and pdf.page_no() > 0: 
                pdf.add_page()
            pdf.add_typology_title(typo_idx, nom_typo)
            for section in typo_data.get("sections", []):
                titre_section = section.get("titre")
                num_chapitre = next((item['number'] for item in chapter_map 
                                   if item['nom_typo'] == nom_typo and item['titre_section'] == titre_section), "")
                texte_genere = typo_previews.get(titre_section)
                if texte_genere:
                    if pdf.get_y() > 250: 
                        pdf.add_page()
                    pdf.add_section_title(num_chapitre, titre_section)
                    
                    # Séparer les parties du texte
                    main_text, placeholders, exemples, crossrefs = extract_parts_for_export(texte_genere, chapter_map, include_ai_notes)
                    
                    # Écrire le texte principal
                    if main_text:
                        pdf.add_body_text(main_text)
                    
                    # Ajouter les éléments supplémentaires s'ils existent et si demandé
                    if include_ai_notes:
                        if placeholders:
                            pdf.add_body_text("\nÀ compléter :")
                            for placeholder in placeholders:
                                pdf.add_body_text(f"• {placeholder}")
                        
                        if exemples:
                            pdf.add_body_text("\nExemples :")
                            for exemple in exemples:
                                pdf.add_body_text(f"• {exemple}")
                    
                    if crossrefs:
                        pdf.add_body_text("\nRéférences croisées :")
                        for crossref in crossrefs:
                            pdf.add_body_text(f"• {crossref}")
        pdf.output(pdf_path)
        return send_file(pdf_path, as_attachment=True, download_name=f"CCTP_{datetime.now().strftime('%Y%m%d')}.pdf")
    except Exception as e:
        print(f"Erreur d'export PDF: {e}")
        return jsonify({"error": f"Erreur lors de la génération du PDF: {str(e)}"}), 500

@app.route('/api/export/docx', methods=['POST'])
def export_word():
    """Génère un fichier Word et le renvoie pour téléchargement."""
    if not DOCX_AVAILABLE:
        return jsonify({"error": "La librairie pour l'export Word (python-docx) n'est pas installée sur le serveur."}), 501
    
    data = request.json
    project_data = data.get('project', [])
    previews_data = data.get('previews', {})
    chapter_map = data.get('chapterMap', [])
    include_ai_notes = data.get('includeAiNotes', True)  # Par défaut, inclure les notes IA

    try:
        word_path = os.path.join(BASE_DIR, 'cctp_export.docx')
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        placeholder_color = RGBColor(220, 53, 69)
        reference_color = RGBColor(40, 167, 69)
        cross_ref_color = RGBColor(0, 123, 255)
        
        for typo_idx, typo_data in enumerate(project_data):
            nom_typo = typo_data.get("nomTypologie")
            typo_previews = previews_data.get(nom_typo, {})
            if not any(typo_previews.values()): 
                continue
            doc.add_heading(f"{typo_idx + 1}. {nom_typo}", level=1)
            for section in typo_data.get("sections", []):
                titre_section = section.get("titre")
                num_chapitre = next((item['number'] for item in chapter_map 
                                   if item['nom_typo'] == nom_typo and item['titre_section'] == titre_section), "")
                texte_genere = typo_previews.get(titre_section)
                
                if texte_genere:
                    doc.add_heading(f"{num_chapitre} {titre_section}", level=2)
                    
                    # Séparer les parties du texte
                    main_text, placeholders, exemples, crossrefs = extract_parts_for_export(texte_genere, chapter_map, include_ai_notes)
                    
                    # Texte principal
                    if main_text:
                        doc.add_paragraph(main_text)
                    
                    # Ajouter les éléments supplémentaires s'ils existent et si demandé
                    if include_ai_notes:
                        # À compléter
                        if placeholders:
                            doc.add_paragraph("À compléter :")
                            for placeholder in placeholders:
                                doc.add_paragraph(f"• {placeholder}")
                        
                        # Exemples
                        if exemples:
                            doc.add_paragraph("Exemples :")
                            for exemple in exemples:
                                doc.add_paragraph(f"• {exemple}")
                    
                    # Références croisées
                    if crossrefs:
                        doc.add_paragraph("Références croisées :")
                        for crossref in crossrefs:
                            doc.add_paragraph(f"• {crossref}")
        doc.save(word_path)
        return send_file(word_path, as_attachment=True, download_name=f"CCTP_{datetime.now().strftime('%Y%m%d')}.docx")
    except Exception as e:
        print(f"Erreur d'export Word: {e}")
        return jsonify({"error": f"Erreur lors de la génération du Word: {str(e)}"}), 500

@app.route('/api/analyze-pdfs', methods=['POST'])
def analyze_pdfs():
    """Analyse un dossier de PDFs pour créer/mettre à jour la base de connaissances."""
    global analysis_status
    
    if analysis_status["running"]:
        return jsonify({"error": "Une analyse est déjà en cours"}), 400
    
    data = request.json
    pdf_directory = data.get('directory')
    
    if not pdf_directory or not os.path.exists(pdf_directory):
        return jsonify({"error": "Dossier spécifié invalide ou inexistant"}), 400
    
    # Vérifier s'il y a des fichiers PDF dans le dossier
    pdf_files = glob.glob(os.path.join(pdf_directory, "*.pdf"))
    if not pdf_files:
        return jsonify({"error": "Aucun fichier PDF trouvé dans le dossier spécifié"}), 400
    
    # Lancer l'analyse en arrière-plan
    analysis_status = {"running": True, "progress": 0, "max_files": len(pdf_files), "current_file": "", "error": None}
    threading.Thread(target=analyze_pdfs_thread, args=(pdf_directory,), daemon=True).start()
    
    return jsonify({"message": "Analyse des PDFs démarrée", "total_files": len(pdf_files)})

@app.route('/api/analyze-pdfs/status', methods=['GET'])
def get_analysis_status():
    """Récupère le statut de l'analyse en cours."""
    return jsonify(analysis_status)

@app.route('/api/browse-directories', methods=['POST'])
def browse_directories():
    """Navigue dans les dossiers pour sélectionner un dossier de PDFs."""
    try:
        data = request.get_json() if request.is_json else {}
        requested_path = data.get('path')
        
        # Déterminer le chemin à utiliser
        if requested_path and os.path.exists(requested_path):
            current_path = os.path.abspath(requested_path)
        else:
            current_path = os.path.expanduser('~')  # Démarrer dans le dossier utilisateur
        
        # Vérifier que le chemin est valide
        if not os.path.exists(current_path):
            current_path = os.path.expanduser('~')
        
        items = []
        
        # Ajouter le dossier parent si on n'est pas à la racine
        parent_path = os.path.dirname(current_path)
        if current_path != parent_path and parent_path:
            items.append({
                'name': '.. (Dossier parent)',
                'path': parent_path,
                'type': 'parent'
            })
        
        # Lister les dossiers et fichiers
        try:
            for item in sorted(os.listdir(current_path)):
                item_path = os.path.join(current_path, item)
                try:
                    if os.path.isdir(item_path):
                        items.append({
                            'name': item,
                            'path': item_path,
                            'type': 'directory'
                        })
                    elif item.lower().endswith('.pdf'):
                        items.append({
                            'name': item,
                            'path': item_path,
                            'type': 'pdf'
                        })
                except OSError:
                    # Ignorer les fichiers/dossiers inaccessibles
                    continue
        except PermissionError:
            return jsonify({"error": "Accès refusé à ce dossier"}), 403
        except OSError as e:
            return jsonify({"error": f"Erreur d'accès au dossier: {str(e)}"}), 500
        
        return jsonify({
            'current_path': current_path,
            'items': items,
            'pdf_count': len([item for item in items if item['type'] == 'pdf'])
        })
        
    except Exception as e:
        print(f"Erreur dans browse_directories: {e}")
        return jsonify({"error": f"Erreur lors de la navigation: {str(e)}"}), 500

def analyze_pdfs_thread(pdf_directory):
    """Thread d'analyse des PDFs et génération de la base de connaissances JSON."""
    global analysis_status
    
    try:
        # Structure JSON pour stocker les données
        knowledge_base_data = {}
        pdf_files = glob.glob(os.path.join(pdf_directory, "*.pdf"))
        
        for i, pdf_path in enumerate(pdf_files):
            filename = os.path.basename(pdf_path)
            analysis_status["progress"] = i
            analysis_status["current_file"] = filename
            
            try:
                with open(pdf_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                # Traiter le texte et créer les sections avec structure hiérarchique
                if text.strip():
                    sections = _parse_document_structure(text)
                    
                    # Créer une typologie principale avec toutes les sections
                    if sections:
                        knowledge_base_data[filename] = [{
                            "nom_typo": filename.replace('.pdf', '').replace('_', ' ').title(),
                            "sections": sections
                        }]
                
            except Exception as e:
                print(f"AVERTISSEMENT: Impossible de lire le fichier {filename}: {e}")
                continue
        
        # Trier les données par nom de fichier pour un résultat cohérent
        sorted_data = dict(sorted(knowledge_base_data.items()))
        
        # Créer le répertoire knowledge_base s'il n'existe pas
        knowledge_base_dir = os.path.join(os.path.dirname(__file__), "knowledge_base")
        os.makedirs(knowledge_base_dir, exist_ok=True)
        
        # Sauvegarder la base de connaissances au format JSON
        json_path = os.path.join(knowledge_base_dir, "knowledge_base.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(sorted_data, f, indent=2, ensure_ascii=False)
        
        # Également sauvegarder au format texte pour compatibilité avec l'existant
        with open(KNOWLEDGE_BASE_PATH, "w", encoding="utf-8") as f:
            contenu_complet = ""
            for filename, data in sorted_data.items():
                contenu_complet += f"\n\n--- EXTRAIT DU DOCUMENT : {filename} ---\n\n"
                for typo in data:
                    for section in typo["sections"]:
                        contenu_complet += f"{section['titre']}\n{section['contenu']}\n\n"
            f.write(contenu_complet)
        
        analysis_status["progress"] = len(pdf_files)
        analysis_status["current_file"] = "Terminé"
        analysis_status["running"] = False
        
    except Exception as e:
        analysis_status["error"] = str(e)
        analysis_status["running"] = False
        print(f"Erreur lors de l'analyse des PDFs: {e}")

def _parse_document_structure(text):
    """Parse le texte pour extraire la structure hiérarchique des sections."""
    sections = []
    
    # Diviser en lignes pour un traitement plus précis
    lines = text.split('\n')
    
    current_section = None
    current_content = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Pattern pour détecter les titres de section avec numérotation
        # Exemples: "2 DESCRIPTION GENERALE", "2.1 Convention de nomenclature", "2.1.1 Modules"
        section_pattern = r'^\s*(\d+(?:\.\d+)*)\s+(.+)$'
        match = re.match(section_pattern, line)
        
        if match:
            section_number = match.group(1)
            section_title = match.group(2).strip()
            
            # Vérifier s'il y a du contenu après ce titre
            has_content_after = False
            content_lines = []
            
            # Regarder les lignes suivantes pour voir s'il y a du contenu
            for j in range(i + 1, len(lines)):
                next_line = lines[j].strip()
                
                # Si on trouve une ligne vide, continuer
                if not next_line:
                    continue
                
                # Si on trouve un autre titre de section, arrêter
                if re.match(r'^\s*\d+(?:\.\d+)*\s+.+$', next_line):
                    break
                
                # Sinon, c'est du contenu
                has_content_after = True
                break
            # Si on a du contenu après ce titre, c'est une vraie section
            if has_content_after:
                # Sauvegarder la section précédente s'il y en a une
                if current_section and current_content:
                    sections.append({
                        "titre": current_section,
                        "contenu": "\n".join(current_content).strip()
                    })
                
                # Commencer une nouvelle section
                current_section = section_title
                current_content = []
            
        elif current_section:
            # Ajouter la ligne au contenu de la section courante
            if line:  # Ignorer les lignes vides
                current_content.append(line)
    
    # Ajouter la dernière section si elle existe
    if current_section and current_content:
        sections.append({
            "titre": current_section,
            "contenu": "\n".join(current_content).strip()
        })
    
    # Si aucune section structurée n'a été trouvée, essayer l'approche par paragraphes
    if not sections:
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        for paragraph in paragraphs[:20]:  # Limiter à 20 sections
            lines = paragraph.split('\n')
            first_line = lines[0].strip()
            
            # Vérifier si la première ligne ressemble à un titre
            if len(first_line) < 150 and len(lines) > 1:
                sections.append({
                    "titre": first_line,
                    "contenu": "\n".join(lines[1:]).strip()
                })
            else:
                # Utiliser les premiers mots comme titre
                words = first_line.split()[:8]
                title = " ".join(words)
                if len(first_line) > len(title):
                    title += "..."
                
                sections.append({
                    "titre": title,
                    "contenu": paragraph
                })
    
    # Si toujours rien, créer une section par défaut
    if not sections and text.strip():
        sections.append({
            "titre": "Document",
            "contenu": text.strip()
        })
    
    return sections

@app.route('/api/models/<model_name>/typologies/<int:typo_index>/sections/<int:section_index>', methods=['DELETE'])
def delete_section(model_name, typo_index, section_index):
    """Supprime une section d'une typologie dans un modèle."""
    project_path = os.path.join(MODELES_DIR, f"{model_name}.json")
    previews_path = os.path.join(PREVIEWS_DIR, f"{model_name}_previews.json")

    if not os.path.exists(project_path):
        return jsonify({"error": "Modèle non trouvé."}), 404

    try:
        # Charger les données du projet
        with open(project_path, 'r', encoding='utf-8') as f:
            project_data = json.load(f)
        
        # Vérifier que la typologie existe
        if typo_index >= len(project_data):
            return jsonify({"error": "Index de typologie invalide."}), 400
        
        typologie = project_data[typo_index]
        sections = typologie.get('sections', [])
        
        # Vérifier que la section existe
        if section_index >= len(sections):
            return jsonify({"error": "Index de section invalide."}), 400
        
        # Récupérer le titre de la section à supprimer pour nettoyer les previews
        section_titre = sections[section_index].get('titre', '')
        nom_typologie = typologie.get('nomTypologie', '')
        
        # Supprimer la section
        sections.pop(section_index)
        
        # Sauvegarder le projet modifié
        with open(project_path, 'w', encoding='utf-8') as f:
            json.dump(project_data, f, indent=2, ensure_ascii=False)
        
        # Nettoyer les previews associées
        previews_data = {}
        if os.path.exists(previews_path):
            with open(previews_path, 'r', encoding='utf-8') as f:
                previews_data = json.load(f)
        
        # Supprimer la preview de la section supprimée
        if nom_typologie in previews_data and section_titre in previews_data[nom_typologie]:
            del previews_data[nom_typologie][section_titre]
        
        # Sauvegarder les previews modifiées
        with open(previews_path, 'w', encoding='utf-8') as f:
            json.dump(previews_data, f, indent=2, ensure_ascii=False)
        
        # Mettre à jour la bibliothèque de sections
        enrichir_bibliotheque_sections()
        
        return jsonify({
            "message": f"Section '{section_titre}' supprimée avec succès.",
            "project": project_data,
            "previews": previews_data
        })
        
    except Exception as e:
        return jsonify({"error": f"Erreur lors de la suppression de la section: {str(e)}"}), 500

@app.route('/api/improve-prompt', methods=['POST'])
def improve_prompt():
    """Ajoute une instruction supplémentaire au prompt de base de VF_LOW_TOKEN."""
    if not openai.api_key:
        return jsonify({"error": "La clé API OpenAI n'est pas configurée sur le serveur."}), 503

    data = request.json
    modification = data.get('modification', '')

    if not modification:
        return jsonify({"error": "La modification est requise."}), 400

    try:
        # Obtenir le prompt de base de VF_LOW_TOKEN
        base_prompt = _get_base_prompt_generation()
        
        # Ajouter l'instruction supplémentaire au prompt de base
        improved_prompt = _apply_prompt_modifications(base_prompt, modification)
        
        return jsonify({"improvedPrompt": improved_prompt})
    except Exception as e:
        return jsonify({"error": f"Erreur lors de l'amélioration du prompt: {str(e)}"}), 500

# --- Lancement de l'application ---
if __name__ == '__main__':
    # Au démarrage, s'assurer que la bibliothèque de sections est à jour
    enrichir_bibliotheque_sections()
    # Lance le serveur de développement Flask
    # En production, utilisez un serveur WSGI comme Gunicorn ou Waitress
    app.run(host='0.0.0.0', port=5000, debug=True)